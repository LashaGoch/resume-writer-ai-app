import os
from dotenv import load_dotenv
from docx import Document
from docxtpl import DocxTemplate
import json
from crewai import Crew, Agent, Task
from flask import Flask, request, render_template, send_file, Response
from markdown import markdown
import re


# Load environment variables
load_dotenv()

app = Flask(__name__)

# Authentication credentials
USERNAME = "canary"
PASSWORD = "resume2025"

# Authentication function
def check_auth(username, password):
    """Check if a username/password combination is valid."""
    return username == USERNAME and password == PASSWORD

def authenticate():
    """Send a 401 response to prompt for credentials."""
    return Response(
        "Could not verify your access level for that URL.\n"
        "You have to login with proper credentials", 401,
        {"WWW-Authenticate": 'Basic realm="Login Required"'}
    )

@app.before_request
def require_auth():
    """Require authentication for all routes."""
    auth = request.authorization
    if not auth or not check_auth(auth.username, auth.password):
        return authenticate()
    

def extract_text_from_docx(file):
    document = Document(file)
    full_text = []

    # Extract text from regular paragraphs
    for para in document.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    # Extract text from tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    full_text.append(cell_text)

    return "\n".join(full_text)


def clean_json_block(text: str) -> str:
    """
    Clean triple-backtick-wrapped JSON content and remove control characters.
    """
    text = text.strip()
    if text.startswith("```json"):
        text = text[7:].strip("` \n")
    elif text.startswith("```"):
        text = text.strip("` \n")
    # Remove control characters except for \n, \r, \t
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    return text

def pad_list(lst, length, filler=""):
    """Pad or truncate list to the desired length."""
    return lst + [filler] * (length - len(lst)) if len(lst) < length else lst[:length]


def render_new_format(context, template_filename="TraditionalFormat.docx", output_path="Final_Resume.docx"):
    try:
        # Get absolute paths
        base_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(base_dir, 'templates', template_filename)
        output_path = os.path.join(base_dir, output_path)
        
        print(f"Template path: {template_path}")
        print(f"Output path: {output_path}")
        
        doc = DocxTemplate(template_path)
        doc.render(context)
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error in render_new_format: {e}")
        return False

# Add this function after your existing imports
def format_resume_markdown(tasks):
    """Convert AI output to formatted markdown"""
    markdown_text = ""
    title = ""

    # First, extract the title from Job Description Writer
    for task in tasks:
        if hasattr(task, 'output') and getattr(task.agent, "role", "") == "Job Description Writer":
            try:
                raw = (getattr(task.output, 'raw_output', None) 
                       or getattr(task.output, 'value', None) 
                       or str(task.output))
                cleaned = clean_json_block(raw)
                data = json.loads(cleaned)
                if isinstance(data, dict) and "experience" in data and data["experience"]:
                    title = data["experience"][0].get("title", "")
            except Exception as e:
                print(f"Error extracting title: {e}")

    # Now, build the markdown
    for task in tasks:
        if not hasattr(task, 'output'):
            continue

        try:
            raw = (getattr(task.output, 'raw_output', None) 
                   or getattr(task.output, 'value', None) 
                   or str(task.output))
            cleaned = clean_json_block(raw)
            data = json.loads(cleaned)


            if task.agent.role == "Name Generator":
                markdown_text += (
                    f"# {data.get('full_name', '')}\n"
                    f"{data.get('location', '')} • {data.get('phone', '')} • "
                    f"[{data.get('email', '')}](mailto:{data.get('email', '')}) • "
                    f"{data.get('LinkedIn', '')}\n\n"
                )
                # Insert the title right after the name block
                if title:
                    markdown_text += f"## {title}\n\n"
            
            elif task.agent.role == "Keyword Generator":
                markdown_text += "\n"
                t_keywords = data.get('top_keywords', [])
                markdown_text += " • ".join(t_keywords) + "\n\n"
            
            elif task.agent.role == "Summary Writer":
                markdown_text += "## Professional Summary\n"
                for summary in data.get('summaries', []):
                    markdown_text += f"{summary}\n\n"
            
            elif task.agent.role == "Areas of Expertise Writer":
                markdown_text += "## Areas of Expertise\n"
                keywords = data.get('expertise_keywords', [])
                # Create 3x3 grid
                for i in range(0, 9, 3):
                    row = keywords[i:i+3]
                    markdown_text += " • ".join(row) + "\n"
                markdown_text += "\n"
            
            elif task.agent.role == "Achievements Writer":
                markdown_text += "## Notable Achievements\n"
                achievements = data.get('notable_achievements', [])
                for achievement in achievements:
                    label = achievement.get('label', '')
                    text = achievement.get('text', '')
                    markdown_text += f"- **{label}:** {text}\n"
                markdown_text += "\n"
            
            elif task.agent.role == "Job Description Writer":
                markdown_text += "## Professional Experience\n"
                for job in data.get('experience', []):
                    markdown_text += f"### {job.get('company')} – {job.get('location')}\n"
                    markdown_text += f"**{job.get('title')}** • {job.get('dates')}\n"
                    markdown_text += f"\n{job.get('description')}\n\n"
                    for achievement in job.get('achievements', []):
                        markdown_text += f"* **{achievement.get('label')}:** {achievement.get('text')}\n"
                    markdown_text += "\n"
            
            elif task.agent.role == "Additional Experience Writer":
                markdown_text += "## Additional Experience\n"
                for job in data.get('earlier_experience', []):
                    markdown_text += (f"**{job.get('company')}** – {job.get('location')}\n"
                                    f"*{job.get('title')}* • {job.get('dates')}\n\n")
            
            elif task.agent.role == "Education Writer":
                markdown_text += "## Education\n"
                for edu in data.get('education', []):
                    markdown_text += (f"**{edu.get('institution')}** • {edu.get('location')}\n"
                                    f"{edu.get('credential')}\n\n")
            
            elif task.agent.role == "Certifications Writer":
                markdown_text += "## Certifications\n"
                for cert in data.get('certifications', []):
                    markdown_text += f"* {cert.get('credential')} – {cert.get('institution')}\n"
                markdown_text += "\n"
                
        except Exception as e:
            print(f"Error formatting {task.agent.role} output: {e}")
            
    return markdown_text




@app.route('/')
def home():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process_resume():
    uploaded_file = request.files.get('file')
    if not uploaded_file:
        return "No file uploaded", 400

    # Extract text from uploaded resume
    resume_text = extract_text_from_docx(uploaded_file)
        
    # Initialize OpenAI key (needed internally by CrewAI)
    openai_api_key = os.getenv('OPENAI_API_KEY')
    if not openai_api_key:
        return "OpenAI API Key not found!", 500

    # Define CrewAI agents

    # Agent to extract name, contact, location
    name_generator = Agent(
        role="Name Generator",
        goal="Extract structured personal information from a resume in structured JSON format.",
        backstory="An expert at reading resumes and identifying core personal information including name, location, phone number, email, and LinkedIn.",
        model="gpt-4o",
        verbose=True,
        allow_delegation=False
    )

    # Agent to extract ATS-friendly keywords
    keyword_generator = Agent(
        role="Keyword Generator",
        goal="Generate exactly four two-word, ATS-optimized keywords based on resume content, suitable for inclusion beneath the candidate's name on a resume.",
        backstory="You are an expert resume keyword analyst trained in recruiting and Applicant Tracking Systems. You extract four distinct, high-impact, two-word phrases that summarize a candidate’s professional strengths and focus areas, tailored to the job title and experience level.",
        model="gpt-4o",
        verbose=True,
        allow_delegation=False
    )

    summary_writer = Agent(
        role="Summary Writer",
        goal="Generate a concise 3-paragraph professional summary from resume content using a consistent, formulaic structure.",
        backstory=(
        "You are an expert in crafting ATS-optimized professional summaries that present candidates with clarity, structure, and strategic positioning. "
        "Each summary must use a predefined 3-paragraph structure with consistent sentence patterns and word choice, reflecting the candidate's experience, communication strengths, and forward-looking value."  
        ),
        model="gpt-4o",
        verbose=True,
        allow_delegation=False
    )

    expertise_writer = Agent(
        role="Areas of Expertise Writer",
        goal="Generate 9 expertise keywords in 3x3 format without repeating top ATS keywords in structured JSON format.",
        backstory="An expert in resume writing and applicant tracking systems. Selects precise, two-word industry phrases that reflect a candidate’s most relevant weekly-used strengths, avoiding redundancy with other keyword sections.",
        model="gpt-4o",
        verbose=True,
        allow_delegation=False
    )

    achievement_writer = Agent(
        role="Achievements Writer",
        goal="Craft 3–5 professional achievement bullets with measurable outcomes using strong verbs and domain-specific bolded keywords.",
        backstory=(
            "You are an expert in resume writing, specializing in turning work experience into high-impact, quantifiable bullet points. "
            "You understand resume tone, industry nuance, and how to highlight both leadership and collaborative achievements with precision. You always follow strict formatting rules."
        ),
        model="gpt-4o",
        verbose=True,
        allow_delegation=False
    )


    experience_writer = Agent(
        role="Job Description Writer",
        goal="Craft concise, structured job experience entries with professional responsibilities and labeled achievements from resume data.",
        backstory="An expert resume editor that produces ATS-optimized job experience sections. Accurately summarizes core responsibilities in a 3-sentence paragraph and highlights up to four labeled, metric-based achievements.",
        model="gpt-4o",
        verbose=True,
        allow_delegation=False
    )

    additional_exp_writer = Agent(
        role="Additional Experience Writer",
        goal="Return earlier work experience entries in structured JSON format.",
        backstory="Summarizes older work experience in a clean, compact structure including company, location, title, and dates.",
        model="gpt-4o",
        verbose=True,
        allow_delegation=False
    )


    education_writer = Agent(
        role="Education Writer",
        goal="Extract and return education entries in structured JSON format.",
        backstory="An expert in parsing and formatting educational history for professional resumes.",
        model="gpt-4o",
        verbose=True,
        allow_delegation=False
    )

    cert_writer = Agent(
        role="Certifications Writer",
        goal="Extract and return certification entries in structured JSON format.",
        backstory="Specialist in identifying and formatting certifications and credentials from resumes.",
        model="gpt-4o",
        verbose=True,
        allow_delegation=False
    )
 

    # Define tasks
    tasks = [

        Task(
            description=(
                f"Extract the following fields from the resume delimited by < >:\n"
                f"1. Full Name\n"
                f"2. Location (City, State)\n"
                f"3. Phone Number\n"
                f"4. Email Address\n"
                f"5. LinkedIn URL. Omit https://www.\n\n"
                f"Resume:\n<{resume_text}>"
            ),
            agent=name_generator,
            expected_output=(
                "Return a JSON object with the following keys:\n"
                "{\n"
                '  "full_name": "Jasmine Taylor",\n'
                '  "location": "New York, NY",\n'
                '  "phone": "555-123-4567",\n'
                '  "email": "jasmine@example.com",\n'
                '  "LinkedIn": "linkedin.com/in/jasminetaylor"\n'
                "}"
            )
        ),

        Task(
            description=(         
                f"Read the resume below (delimited by < >) and extract the top four (4) two-word keywords optimized for Applicant Tracking Systems (ATS).\n" 
                f"Resume:\n<{resume_text}>\n\n"
                f"OUTPUT RULES:\n"
                f"• No repeated concepts, soft skills, or personal traits.\n"
                f"• Use ampersands (&) only when standard (e.g., Risk & Compliance).\n"
                f"• These should be skills or phrases that match professional strengths and job market terminology.\n"
                f"• Match terms to job level and job description keywords."
                
            ),
            agent=keyword_generator,
            expected_output=(
                "Return a JSON list of exactly 4 strings, like this:\n"
                " {\n"
                ' "top_keywords": [\n'
                '       "keyword 1", "keyword 2", "keyword 3", "keyword 4"\n'
                "   ]\n"
                "}"
            )
        ),

        Task(
            description=(
                f"Read the resume below (delimited by < >) and write a concise 3-paragraph professional summary.\n"
                f"Resume:\n<{resume_text}>\n\n"
                f"Follow this structure:\n"
                f"Each paragraph should follow a three-sentence structure.\n\n"
                f"Resume:\n<{resume_text}>"
                f"Paragraph 1 – Experience & Impact:\n"
                f"{{Descriptor 1}} and {{Descriptor 2}} {{Role Noun}} offering {{Years}}+ years of experience {{Action 1}}, {{Action 2}}, and {{Action 3}} in {{Industry/Function}}.\n\n"
                f"Paragraph 2 – Communication & Influence:\n"
                f"{{Descriptor 1}} and {{Descriptor 2}} {{Role Noun}} skilled at {{Soft Skill A}}, {{Soft Skill B}}, and {{Outcome}} using {{Trait A}}, {{Trait B}}, and {{Trait C}}.\n\n"
                f"Paragraph 3 – Forward Value & Mission:\n"
                f"{{Descriptor 1}} and {{Descriptor 2}} {{Role Noun}} focused on {{Mission A}}, {{Mission B}}, and {{Mission C}} by {{How they do it}}, delivering {{Impact}}.\n\n"
                f"Use random sampling from the following predefined lists:\n"
                f"- Descriptors: strategic, collaborative, detail-oriented, visionary, innovative, adaptable, people-focused, entrepreneurial, composed, solutions-oriented, future-facing\n"
                f"- Nouns (Role): leader, communicator, problem solver, collaborator, expert, strategist, business partner, relationship builder\n"
                f"- Soft Skills / Traits: emotional intelligence, storytelling ability, cultural awareness, growth mindset, calm under pressure, adaptability, strategic insight, hands-on approach\n"
                f"- Missions: improving access, driving sustainability, transforming service delivery\n"
                f"- Impacts: global health, community growth, team cohesion, policy change\n\n"
                f"Each paragraph must be exactly one sentence, 25–30 words."
            ),
            agent=summary_writer,
            expected_output=(
                "Return a JSON object with the following structure:\n\n"
                "{\n"
                '  "summaries": [\n'
                '    "Paragraph 1",\n'
                '    "Paragraph 2",\n'
                '    "Paragraph 3"\n'
                "  ]\n"
                "}\n\n"
                "Ensure each paragraph is complete, ATS-friendly, and aligned with modern resume standards."
            )
        ),

        Task(
            description=(
                f"From the resume below, generate 9 two-word unique 'Areas of Expertise' keyword phrases arranged for visual formatting in 3 columns and 3 rows.\n "
                f"Resume:\n<{resume_text}>\n\n"
                f"Do not repeat the top keywords already used earlier. Keep all keywords concise and resume-appropriate.\n"
                f"All phrases should reflect weekly-used, high-signal competencies.\n"
                f"Maintain balance across technical, strategic, and operational skills.\n"
                f"Only use two-word phrases that would appear in real job descriptions or LinkedIn.\n"
            
            ),
            agent=expertise_writer,
            expected_output=(
                "Return a JSON object like this:\n"
                '{\n'
                '  "expertise_keywords": [\n'
                '    "Keyword 1", "Keyword 2", "Keyword 3",\n'
                '    "Keyword 4", "Keyword 5", "Keyword 6",\n'
                '    "Keyword 7", "Keyword 8", "Keyword 9"\n'
                '  ]\n'
                '}\n\n'
                "Only include keyword phrases. No explanations, no formatting, no column titles."
            )
        ),

        Task(
            description=(
                f"Based on the resume content below (delimited by < >), write 3–5 bullet points describing notable professional achievements. "
                f"Resume:\n<{resume_text}>\n\n"
                f"Each bullet must be measurable, impactful, exactly 1 sentence, 25–35 words long.\n\n"
                f"Start with **two bolded keywords** describing the domain/theme (e.g., **Revenue Growth:**)\n"
                f"Begin the sentence with a strong action verb (e.g., Spearheaded, Delivered, Improved)\n"
                f"Include real, measurable outcomes (no fake metrics)\n"
                f"Avoid all pronouns, articles, adverbs, and passive voice\n"
                f"Use direct/assertive verbs if the user led the effort; use collaborative framing if it was a team effort\n"
                f"If 'and' is used more than twice, replace the third instance with 'as well as', 'in addition to', etc.\n"
                f"Do not invent metrics for achievements. Only include information explicitly stated in the resume text provided.\n" 
                f"If the achievement does not have an outcome, you can infer (without making up metrics) what the likely outcome was.\n"
                f"Tailor each bullet to the resume content and role."
            ),
            agent=achievement_writer,
            expected_output=(
                "Return a JSON object with this structure:\n"
                "{\n"
                '  "notable_achievements": [\n'
                '        {"label": "Lead Generation", "text": "" Improved lead generation by 150% across nine websites through the deployment of AI and A/B testing tools, resulting in four successful campaign implementations within six months."},\n'
                '        {"label": "Reporting Centralization", "text": " Led the centralization of marketing reporting across the organization by integrating six data sources into Qlik, generating insights for 33+ KPIs and enhancing decision-making for over 300 users."}\n'
                '        {"label": "Database Growth", "text": " Increased prospects and customers database from 100,000 to 800,000 within one year by utilizing both external and internal channels, significantly boosting marketing outreach and engagement."}\n'
                "  ]\n"
                "}\n\n"
                "Only return the JSON. Do not include commentary, headers, or formatting."
            )
        ),

        Task(
            description=(
                f"Use the resume below (delimited by < >) to write structured job descriptions for the 2–3 most recent roles.\n"
                f"Resume:\n<{resume_text}>\n"
                f"For each role, return:\n"
                f"• Company name\n"
                f"• Location (City, State or Country)\n"
                f"• Title\n"
                f"• Dates of employment\n"
                f"• 3-sentence paragraph (≤50 words) describing core job responsibilities ONLY — no achievements, no metrics\n"
                f"• 1–4 achievement bullet points using this structure:\n"
                f"    {{\"label\": \"Verb\", \"text\": \"achievement text with outcome/metric\"}}\n\n"
                f"DESCRIPTION RULES:\n"
                f"• Begin with a high-level task summary\n"
                f"• Next two sentences describe recurring responsibilities using action verbs\n"
                f"ACHIEVEMENT RULES:\n"
                f"• 1–4 bullets per role\n"
                f"• Each bullet starts with a label verb (e.g., \"Led\", \"Drove\", \"Built\")\n"
                f"• Text must contain a measurable result (metric, percentage, impact)\n"
                f"• Use short, strong phrasing\n"
                f"• Avoid repeating notable achievements that were already listed earlier\n"
            ),
            agent=experience_writer,
            expected_output=(
                "Return a JSON object with the following structure:\n\n"
                "{\n"
                '  "experience": [\n'
                "    {\n"
                '      "company": "Company Name",\n'
                '      "location": "City, State",\n'
                '      "title": "Job Title",\n'
                '      "dates": "Start Year-End Year or Present",\n'
                '      "description": "Three-sentence responsibility paragraph here (≤50 words).",\n'
                '      "achievements": [\n'
                '        {"label": "Led", "text": "achievement with metric."},\n'
                '        ...\n'
                '      ]\n'
                "    },\n"
                "    ...\n"
                "  ]\n"
                "}"
            )
        ),

        Task(
            description=(
                f"From the resume below (delimited by < >), extract earlier/older work experience entries (not among the most recent 2–3 roles). "
                f"For each job, return:\n"
                f"- Company name\n"
                f"- Location (City, State)\n"
                f"- Job Title\n"
                f"- Dates of Employment\n\n"
                f"Resume:\n<{resume_text}>"
            ),
            agent=additional_exp_writer,
            expected_output=(
                "Return a JSON object in the following structure:\n\n"
                "{\n"
                '  "earlier_experience": [\n'
                "    {\n"
                '      "company": "Duke University",\n'
                '      "location": "Durham, North Carolina",\n'
                '      "title": "Project Coordinator (Time-Limited Grant)",\n'
                '      "dates": "Jan 2011 – Jan 2012"\n'
                "    },\n"
                "    ...\n"
                "  ]\n"
                "}\n\n"
                "Do not include responsibilities, bullet points, or extra commentary — just the structured entries."
            )
        ),

        Task(
            description=(
                f"From the resume below (delimited by < >), extract all education entries. "
                f"For each entry, return:\n"
                f"- Institution name\n"
                f"- Location (City, State or Online)\n"
                f"- Credential (e.g. degree, diploma)\n\n"
                f"Resume:\n<{resume_text}>"
            ),
            agent=education_writer,
            expected_output=(
                "Return a JSON object with this structure:\n\n"
                "{\n"
                '  "education": [\n'
                "    {\n"
                '      "institution": "Harvard University",\n'
                '      "location": "Cambridge, MA",\n'
                '      "credential": "Master of Business Administration"\n'
                "    },\n"
                "    ...\n"
                "  ]\n"
                "}"
            )
        ),

        Task(
            description=(
                f"From the resume below (delimited by < >), extract all professional certifications or credentials. "
                f"For each entry, return:\n"
                f"- Institution (e.g. issuing organization)\n"
                f"- Location (e.g. Online or city)\n"
                f"- Credential (e.g. certificate title)\n\n"
                f"Resume:\n<{resume_text}>"
            ),
            agent=cert_writer,
            expected_output=(
                "Return a JSON object with this structure:\n\n"
                "{\n"
                '  "certifications": [\n'
                "    {\n"
                '      "institution": "Project Management Institute",\n'
                '      "location": "Online",\n'
                '      "credential": "Project Management Professional (PMP)"\n'
                "    },\n"
                "    ...\n"
                "  ]\n"
                "}"
            )
        )
    ]


    # Run the crew
    crew = Crew(agents=[name_generator, keyword_generator, summary_writer, expertise_writer, 
        achievement_writer, experience_writer, additional_exp_writer, education_writer, cert_writer
    ], tasks=tasks, verbose=True)
    
    result = crew.kickoff()

    compiled_resume_text = ""
    for task in crew.tasks:
        if hasattr(task, 'output'):
            compiled_resume_text += f"\n\n{task.output}"

    # Clean up the compiled resume text
    try:
        # Load template
        template_path = os.path.join(os.path.dirname(__file__), 'templates', 'TraditionalFormat.docx')
        doc = DocxTemplate("templates/TraditionalFormat.docx")
        context = {}

        # Extract Data from Crew Tasks
        for task in crew.tasks:
            if hasattr(task, 'output') and task.output:
                raw = (
                    getattr(task.output, 'raw_output', None)
                    or getattr(task.output, 'value', None)
                    or str(task.output)
                )
                cleaned = clean_json_block(raw)

                try:
                    parsed = json.loads(cleaned)
                    if isinstance(parsed, dict):
                        context.update(parsed)
                    elif isinstance(parsed, list) and task.agent.role == "Keyword Generator":
                        context["top_keywords"] = parsed
                except Exception as e:
                    print(f"❌ Could not parse output from {task.agent.role}:\n{cleaned[:300]}\nError: {e}")

 
        # Optional sections
        for key in ["earlier_experience", "education", "certifications"]:
            context[key] = context.get(key, [])

        # Final Validation and Render
        required_keys = ["experience", "earlier_experience", "education", "certifications"]
        missing = [k for k in required_keys if k not in context]
        if missing:
            print(f"⚠️ Missing fields in context: {missing}")
        else:
            output_path = os.path.join(os.path.dirname(__file__), 'Final_Resume.docx')
            doc.render(context)
            doc.save(output_path)
            print("✅ Resume rendered and saved as Final_Resume.docx")

    except Exception as e:
        print(f"Error processing template: {e}")


    # Return the template as before
    compiled_resume_html = markdown(format_resume_markdown(crew.tasks))
    return render_template('result.html', compiled_resume_html=compiled_resume_html)




@app.route('/download_new_format')
def download_new_format():
    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        file_path = os.path.join(base_dir, "Final_Resume.docx")
        
        if not os.path.exists(file_path):
            return "Resume file not found. Please process your resume first.", 404
            
        return send_file(
            file_path,
            as_attachment=True,
            download_name="Final_Resume.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        print(f"Error downloading file: {e}")
        return f"Error downloading file: {str(e)}", 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080)
